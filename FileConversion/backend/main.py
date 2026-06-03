import io
import logging
import os
import shutil
import subprocess
import tempfile
import uuid
from pathlib import Path

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse

from config import CATEGORIES, FFMPEG_PATH, HOST, PORT, TEMP_DIR

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="File Conversion API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://localhost:4173",
        "https://file-convert-b5a76.firebaseapp.com",
        "https://file-convert-b5a76.web.app",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.makedirs(TEMP_DIR, exist_ok=True)

FORMAT_CATEGORY_MAP = {}
for cat_key, cat_info in CATEGORIES.items():
    for fmt in cat_info["formats"]:
        FORMAT_CATEGORY_MAP[fmt] = cat_key

MIME_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "odt": "application/vnd.oasis.opendocument.text",
    "ods": "application/vnd.oasis.opendocument.spreadsheet",
    "odp": "application/vnd.oasis.opendocument.presentation",
    "rtf": "application/rtf",
    "csv": "text/csv",
    "md": "text/markdown",
    "html": "text/html",
    "xml": "application/xml",
    "mp3": "audio/mpeg",
    "wav": "audio/wav",
    "aac": "audio/aac",
    "flac": "audio/flac",
    "ogg": "audio/ogg",
    "m4a": "audio/mp4",
    "opus": "audio/opus",
    "mp4": "video/mp4",
    "avi": "video/x-msvideo",
    "mkv": "video/x-matroska",
    "mov": "video/quicktime",
    "wmv": "video/x-ms-wmv",
    "flv": "video/x-flv",
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "bmp": "image/bmp",
    "svg": "image/svg+xml",
}


def get_category(fmt: str) -> str | None:
    return FORMAT_CATEGORY_MAP.get(fmt.lower())


def cleanup(paths: list[str]):
    for p in paths:
        try:
            if os.path.isfile(p):
                os.unlink(p)
            elif os.path.isdir(p):
                shutil.rmtree(p, ignore_errors=True)
        except Exception:
            pass


# ─── Image conversion ────────────────────────────────────────────────

def convert_image(input_path: str, output_path: str, target_fmt: str):
    from PIL import Image

    img = Image.open(input_path)
    if target_fmt == "jpg":
        target_fmt = "jpeg"
    if img.mode in ("P", "RGBA") and target_fmt in ("jpeg", "jpg"):
        img = img.convert("RGB")
    img.save(output_path, format=target_fmt.upper())


# ─── Document conversion ─────────────────────────────────────────────

def try_pypandoc(src: str, dst: str, src_fmt: str | None, dst_fmt: str):
    try:
        import pypandoc

        extra_args = []
        if dst_fmt == "pdf":
            extra_args = ["--pdf-engine=xelatex"]
        pypandoc.convert_file(
            src,
            dst_fmt,
            format=src_fmt,
            outputfile=dst,
            extra_args=extra_args,
        )
        return True
    except Exception as e:
        logger.warning("pypandoc failed: %s", e)
        return False


def convert_txt_to_docx(src: str, dst: str):
    from docx import Document

    doc = Document()
    with open(src, encoding="utf-8", errors="replace") as f:
        doc.add_paragraph(f.read())
    doc.save(dst)


def convert_docx_to_txt(src: str, dst: str):
    from docx import Document

    doc = Document(src)
    with open(dst, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")


def convert_txt_to_pdf(src: str, dst: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(dst, pagesize=A4)
    width, height = A4
    y = height - 40
    with open(src, encoding="utf-8", errors="replace") as f:
        for line in f:
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, line.rstrip())
            y -= 14
    c.save()


def convert_md_to_html(src: str, dst: str):
    import markdown

    with open(src, encoding="utf-8") as f:
        html = markdown.markdown(f.read(), extensions=["fenced_code", "tables"])
    with open(dst, "w", encoding="utf-8") as f:
        f.write(html)


def convert_html_to_md(src: str, dst: str):
    import html2text

    with open(src, encoding="utf-8") as f:
        md = html2text.HTML2Text().handle(f.read())
    with open(dst, "w", encoding="utf-8") as f:
        f.write(md)


def convert_csv_to_xlsx(src: str, dst: str):
    import csv

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    with open(src, encoding="utf-8-sig") as f:
        for row in csv.reader(f):
            ws.append(row)
    wb.save(dst)


def convert_xlsx_to_csv(src: str, dst: str):
    import csv

    from openpyxl import load_workbook

    wb = load_workbook(src, read_only=True)
    ws = wb.active
    with open(dst, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(row)
    wb.close()


def convert_pdf_to_txt(src: str, dst: str):
    import fitz

    doc = fitz.open(src)
    text = "\n".join(page.get_text() for page in doc)
    with open(dst, "w", encoding="utf-8") as f:
        f.write(text)
    doc.close()


def convert_document(input_path: str, output_path: str, src_fmt: str, dst_fmt: str):
    if src_fmt == dst_fmt:
        shutil.copy2(input_path, output_path)
        return

    if try_pypandoc(input_path, output_path, src_fmt, dst_fmt):
        return

    pair = (src_fmt, dst_fmt)
    handlers = {
        ("txt", "docx"): convert_txt_to_docx,
        ("docx", "txt"): convert_docx_to_txt,
        ("txt", "pdf"): convert_txt_to_pdf,
        ("md", "html"): convert_md_to_html,
        ("html", "md"): convert_html_to_md,
        ("csv", "xlsx"): convert_csv_to_xlsx,
        ("xlsx", "csv"): convert_xlsx_to_csv,
        ("pdf", "txt"): convert_pdf_to_txt,
    }

    handler = handlers.get(pair)
    if handler:
        handler(input_path, output_path)
        return

    raise HTTPException(
        status_code=400,
        detail=f"不支援的轉換: {src_fmt} → {dst_fmt}。請安裝 pandoc (https://pandoc.org) 以獲得更多格式支援。",
    )


# ─── Audio / Video conversion (FFmpeg) ───────────────────────────────

def convert_media(input_path: str, output_path: str, src_fmt: str, dst_fmt: str):
    if src_fmt == dst_fmt:
        shutil.copy2(input_path, output_path)
        return

    if not shutil.which(FFMPEG_PATH) and not os.path.isfile(FFMPEG_PATH):
        raise HTTPException(
            status_code=400,
            detail=f"找不到 FFmpeg，請安裝 FFmpeg (https://ffmpeg.org) 才能進行音訊/影片轉換。",
        )

    acodec_map = {
        "mp3": "libmp3lame",
        "aac": "aac",
        "flac": "flac",
        "ogg": "libvorbis",
        "opus": "libopus",
    }

    vcodec_map = {
        "mp4": "libx264",
        "mkv": "libx264",
        "mov": "libx264",
        "avi": "libxvid",
        "wmv": "msmpeg4v2",
        "flv": "flv",
    }

    is_same_category = get_category(src_fmt) == get_category(dst_fmt)
    if not is_same_category:
        raise HTTPException(
            status_code=400,
            detail=f"無法在音訊與影片之間轉換: {src_fmt} → {dst_fmt}",
        )

    cat = get_category(src_fmt)
    cmd = [FFMPEG_PATH, "-y", "-i", input_path]

    if cat == "audio":
        acodec = acodec_map.get(dst_fmt)
        if acodec:
            cmd += ["-acodec", acodec]
        cmd += ["-vn", output_path]
    elif cat == "video":
        vcodec = vcodec_map.get(dst_fmt, "libx264")
        cmd += ["-vcodec", vcodec, "-acodec", "aac", output_path]

    result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
    if result.returncode != 0:
        logger.error("FFmpeg error: %s", result.stderr)
        raise HTTPException(status_code=500, detail=f"FFmpeg 轉換失敗: {result.stderr[:200]}")


# ─── API routes ──────────────────────────────────────────────────────

@app.get("/api/formats")
def list_formats():
    return {
        "categories": [
            {"key": k, "label": v["label"], "formats": v["formats"]}
            for k, v in CATEGORIES.items()
        ]
    }


@app.post("/api/convert")
async def convert_file(
    file: UploadFile = File(...),
    target_format: str = Form(...),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="未提供檔案")

    original_name = file.filename
    ext = original_name.rsplit(".", 1)[-1].lower() if "." in original_name else ""
    if not ext:
        raise HTTPException(status_code=400, detail="無法辨識檔案格式")

    target = target_format.lower().lstrip(".")
    src_cat = get_category(ext)
    dst_cat = get_category(target)

    if not src_cat:
        raise HTTPException(status_code=400, detail=f"不支援的來源格式: {ext}")
    if not dst_cat:
        raise HTTPException(status_code=400, detail=f"不支援的目標格式: {target}")
    if src_cat != dst_cat:
        raise HTTPException(
            status_code=400,
            detail=f"無法跨類別轉換: {ext}({src_cat}) → {target}({dst_cat})",
        )

    job_id = uuid.uuid4().hex
    work_dir = os.path.join(TEMP_DIR, job_id)
    os.makedirs(work_dir, exist_ok=True)

    input_path = os.path.join(work_dir, f"input.{ext}")
    output_filename = original_name.rsplit(".", 1)[0] + "." + target
    output_path = os.path.join(work_dir, f"output.{target}")

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        if src_cat == "photo":
            convert_image(input_path, output_path, target)
        elif src_cat == "document":
            convert_document(input_path, output_path, ext, target)
        elif src_cat in ("audio", "video"):
            convert_media(input_path, output_path, ext, target)

        if not os.path.isfile(output_path):
            raise HTTPException(status_code=500, detail="轉換失敗，未產生輸出檔案")

        mime = MIME_TYPES.get(target, "application/octet-stream")

        return FileResponse(
            output_path,
            media_type=mime,
            filename=output_filename,
            headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Conversion failed")
        raise HTTPException(status_code=500, detail=f"轉換失敗: {str(e)}")
    finally:
        cleanup([work_dir])


@app.get("/api/health")
def health_check():
    import time
    return {"status": "ok", "timestamp": time.time(), "service": "file-conversion"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host=HOST, port=PORT, reload=True)
